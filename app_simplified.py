import streamlit as st
import subprocess
import json
from datetime import datetime
import os
from pathlib import Path
import tempfile
import hashlib
import random
import time

# PM Wisdom Quotes for loading states
PM_QUOTES = [
    "The best PMs are great at saying 'no' to good ideas.",
    "Fall in love with the problem, not your solution.",
    "Metrics are lagging indicators. Talk to users.",
    "Ship early, ship often, iterate always.",
    "Your first solution is rarely the right solution.",
    "Perfect is the enemy of shipped.",
    "The best spec is the one that gets built.",
    "User pain × frequency = Priority",
    "Build less, learn more.",
    "Strong opinions, loosely held.",
]

# Custom CSS - Modern Linear/Stripe-inspired design
CUSTOM_CSS = """
<style>
    /* Import Inter font for modern typography */
    @import url('https://rsms.me/inter/inter.css');

    /* Root variables - Linear/Stripe color palette */
    :root {
        --color-bg: #ffffff;
        --color-bg-subtle: #fafafa;
        --color-bg-muted: #f5f5f5;
        --color-border: #e5e5e5;
        --color-border-strong: #d4d4d8;
        --color-text: #18181b;
        --color-text-muted: #71717a;
        --color-text-subtle: #a1a1aa;
        --color-accent: #5b5bd6;
        --color-accent-hover: #4a4ab8;
        --color-accent-light: #ededfc;
        --color-success: #16a34a;
        --color-warning: #ea580c;
        --color-error: #dc2626;
        --shadow-sm: 0 1px 2px 0 rgba(0, 0, 0, 0.05);
        --shadow-md: 0 4px 6px -1px rgba(0, 0, 0, 0.1);
        --shadow-lg: 0 10px 15px -3px rgba(0, 0, 0, 0.1);
        --radius-sm: 6px;
        --radius-md: 8px;
        --radius-lg: 12px;
        --spacing-xs: 4px;
        --spacing-sm: 8px;
        --spacing-md: 16px;
        --spacing-lg: 24px;
        --spacing-xl: 32px;
        --spacing-2xl: 48px;
    }

    /* Global resets and base styles */
    * {
        font-family: 'Inter', -apple-system, BlinkMacSystemFont, 'Segoe UI', sans-serif;
    }

    /* Main container - clean, spacious */
    .main {
        background: var(--color-bg);
        max-width: 800px;
        margin: 0 auto;
        padding: var(--spacing-2xl) var(--spacing-lg);
    }

    .main .block-container {
        padding: 0;
        max-width: 100%;
    }

    /* Typography hierarchy */
    h1 {
        font-size: 32px;
        font-weight: 700;
        color: var(--color-text);
        letter-spacing: -0.02em;
        margin-bottom: var(--spacing-sm);
        line-height: 1.2;
    }

    h2 {
        font-size: 24px;
        font-weight: 600;
        color: var(--color-text);
        letter-spacing: -0.01em;
        margin-top: var(--spacing-xl);
        margin-bottom: var(--spacing-md);
    }

    h3 {
        font-size: 18px;
        font-weight: 600;
        color: var(--color-text);
        margin-top: var(--spacing-lg);
        margin-bottom: var(--spacing-sm);
    }

    p {
        font-size: 15px;
        line-height: 1.6;
        color: var(--color-text-muted);
    }

    /* Remove Streamlit branding and clutter */
    #MainMenu {visibility: hidden;}
    footer {visibility: hidden;}
    header {visibility: hidden;}

    /* Buttons - Linear style with better contrast */
    .stButton > button {
        background: var(--color-accent);
        color: white !important;
        border: none;
        border-radius: var(--radius-md);
        padding: 12px 24px;
        font-size: 15px;
        font-weight: 600;
        transition: all 0.15s ease;
        box-shadow: var(--shadow-sm);
        cursor: pointer;
    }

    .stButton > button:hover {
        background: var(--color-accent-hover);
        box-shadow: var(--shadow-md);
        transform: translateY(-1px);
    }

    .stButton > button:active {
        transform: translateY(0);
    }

    /* Primary button - high contrast */
    .stButton > button[kind="primary"] {
        background: var(--color-text);
        color: white !important;
        font-weight: 600;
    }

    .stButton > button[kind="primary"]:hover {
        background: #000000;
    }

    /* Secondary button variant */
    .stButton > button[kind="secondary"] {
        background: var(--color-bg-subtle);
        color: var(--color-text);
        border: 1px solid var(--color-border);
    }

    .stButton > button[kind="secondary"]:hover {
        background: var(--color-bg-muted);
        border-color: var(--color-border-strong);
    }

    /* Text inputs - clean and minimal */
    .stTextInput input, .stTextArea textarea {
        border: 1px solid var(--color-border);
        border-radius: var(--radius-md);
        padding: 10px 12px;
        font-size: 14px;
        color: var(--color-text);
        background: var(--color-bg);
        transition: all 0.15s ease;
    }

    .stTextInput input:focus, .stTextArea textarea:focus {
        border-color: var(--color-accent);
        box-shadow: 0 0 0 3px var(--color-accent-light);
        outline: none;
    }

    /* Radio buttons - custom styled */
    .stRadio > div {
        gap: var(--spacing-sm);
    }

    .stRadio > div > label {
        background: var(--color-bg);
        border: 1px solid var(--color-border);
        border-radius: var(--radius-md);
        padding: var(--spacing-md);
        cursor: pointer;
        transition: all 0.15s ease;
        display: flex;
        align-items: center;
    }

    .stRadio > div > label:hover {
        border-color: var(--color-accent);
        background: var(--color-accent-light);
    }

    .stRadio > div > label[data-checked="true"] {
        border-color: var(--color-accent);
        background: var(--color-accent-light);
    }

    /* Score card - minimal and elegant */
    .score-card {
        background: var(--color-bg);
        border: 1px solid var(--color-border);
        border-radius: var(--radius-lg);
        padding: var(--spacing-xl);
        text-align: center;
        margin: var(--spacing-lg) 0;
        box-shadow: var(--shadow-sm);
    }

    .score-number {
        font-size: 72px;
        font-weight: 700;
        color: var(--color-text);
        letter-spacing: -0.03em;
        line-height: 1;
        margin: var(--spacing-sm) 0;
    }

    .score-label {
        font-size: 13px;
        font-weight: 500;
        color: var(--color-text-subtle);
        text-transform: uppercase;
        letter-spacing: 0.05em;
    }

    /* Dimension cards - clean list style */
    .dimension-card {
        background: var(--color-bg);
        border: 1px solid var(--color-border);
        border-radius: var(--radius-md);
        padding: var(--spacing-md);
        margin: var(--spacing-sm) 0;
        transition: all 0.15s ease;
    }

    .dimension-card:hover {
        border-color: var(--color-border-strong);
        box-shadow: var(--shadow-sm);
    }

    /* Chat messages - Linear-style conversation */
    .stChatMessage {
        background: var(--color-bg);
        border: 1px solid var(--color-border);
        border-radius: var(--radius-lg);
        padding: var(--spacing-lg);
        margin: var(--spacing-md) 0;
        box-shadow: var(--shadow-sm);
    }

    /* Topic badge - subtle and minimal */
    .topic-badge {
        display: inline-flex;
        align-items: center;
        background: var(--color-bg-muted);
        color: var(--color-text-muted);
        padding: 4px 10px;
        border-radius: 999px;
        font-size: 12px;
        font-weight: 500;
        margin-bottom: var(--spacing-sm);
    }

    /* Progress bar - Linear style */
    .progress-container {
        width: 100%;
        height: 2px;
        background: var(--color-border);
        border-radius: 999px;
        overflow: hidden;
        margin: var(--spacing-lg) 0;
    }

    .progress-bar {
        height: 100%;
        background: var(--color-accent);
        transition: width 0.3s ease;
    }

    .progress-text {
        font-size: 13px;
        color: var(--color-text-subtle);
        text-align: center;
        margin-top: var(--spacing-sm);
        font-weight: 500;
    }

    /* Loading state - minimal spinner */
    .loading-container {
        display: flex;
        flex-direction: column;
        align-items: center;
        justify-content: center;
        padding: var(--spacing-2xl) 0;
    }

    .spinner {
        width: 24px;
        height: 24px;
        border: 2px solid var(--color-border);
        border-top-color: var(--color-accent);
        border-radius: 50%;
        animation: spin 0.8s linear infinite;
    }

    @keyframes spin {
        to { transform: rotate(360deg); }
    }

    .pm-quote {
        font-size: 14px;
        font-style: italic;
        color: var(--color-text-subtle);
        text-align: center;
        margin-top: var(--spacing-md);
        max-width: 400px;
    }

    /* Info/Alert boxes - subtle */
    .stAlert {
        background: var(--color-bg-subtle);
        border: 1px solid var(--color-border);
        border-radius: var(--radius-md);
        padding: var(--spacing-md);
        color: var(--color-text-muted);
        font-size: 14px;
    }

    /* Dividers */
    hr {
        border: none;
        border-top: 1px solid var(--color-border);
        margin: var(--spacing-xl) 0;
    }

    /* Expander - clean accordion */
    .streamlit-expanderHeader {
        background: var(--color-bg);
        border: 1px solid var(--color-border);
        border-radius: var(--radius-md);
        padding: var(--spacing-md);
        font-size: 14px;
        font-weight: 500;
        color: var(--color-text);
    }

    /* Metrics - simple number display */
    .stMetric {
        background: var(--color-bg);
        padding: 0;
    }

    /* Sidebar - minimal */
    .css-1d391kg {
        background: var(--color-bg-subtle);
        border-right: 1px solid var(--color-border);
    }

    /* Question response area */
    .response-card {
        background: var(--color-bg);
        border: 1px solid var(--color-border);
        border-radius: var(--radius-lg);
        padding: var(--spacing-lg);
        margin: var(--spacing-lg) 0;
    }

    .question-display {
        background: var(--color-bg-subtle);
        border-left: 3px solid var(--color-accent);
        border-radius: var(--radius-md);
        padding: var(--spacing-lg);
        margin: var(--spacing-md) 0;
        font-size: 15px;
        line-height: 1.6;
        color: var(--color-text);
    }

    /* Success celebration - subtle */
    .celebration {
        text-align: center;
        padding: var(--spacing-xl) 0;
    }

    .celebration-icon {
        font-size: 48px;
        margin-bottom: var(--spacing-md);
    }

    .celebration-title {
        font-size: 24px;
        font-weight: 600;
        color: var(--color-text);
        margin-bottom: var(--spacing-sm);
    }

    .celebration-subtitle {
        font-size: 14px;
        color: var(--color-text-muted);
    }

    /* Smooth transitions */
    * {
        transition: background-color 0.15s ease, border-color 0.15s ease;
    }
</style>
"""

# Check for python-docx
try:
    from docx import Document
    import io
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False

# Page config
st.set_page_config(
    page_title="PM Spec Analyzer - Socratic Coach",
    page_icon="🎤",
    layout="wide",
    initial_sidebar_state="expanded"
)

# Apply custom CSS
st.markdown(CUSTOM_CSS, unsafe_allow_html=True)

def show_loading_with_quote(message="Thinking..."):
    """Show loading spinner with random PM quote - minimal Linear style"""
    quote = random.choice(PM_QUOTES)

    html = f"""
    <div class="loading-container">
        <div class="spinner"></div>
        <p style="margin-top: 16px; font-size: 14px; font-weight: 500; color: var(--color-text);">{message}</p>
        <p class="pm-quote">"{quote}"</p>
    </div>
    """
    return st.markdown(html, unsafe_allow_html=True)

def show_progress(current, total):
    """Show minimal progress bar"""
    percentage = (current / total) * 100 if total > 0 else 0
    html = f"""
    <div class="progress-container">
        <div class="progress-bar" style="width: {percentage}%"></div>
    </div>
    <p class="progress-text">Question {current} of ~{total}</p>
    """
    st.markdown(html, unsafe_allow_html=True)

# Find Claude Code CLI
def find_claude_cli():
    """Find Claude Code CLI in common locations"""
    claude_paths = [
        Path.home() / ".local/bin/claude",
        Path("/usr/local/bin/claude"),
        Path("/opt/homebrew/bin/claude"),
        Path("/usr/bin/claude"),
    ]

    for path in claude_paths:
        if path.exists() and path.is_file():
            return str(path)

    try:
        result = subprocess.run(['which', 'claude'], capture_output=True, text=True)
        if result.returncode == 0:
            return result.stdout.strip()
    except:
        pass

    return None

CLAUDE_CLI = find_claude_cli()

# Set up local storage directory
STORAGE_DIR = Path.home() / ".pm_spec_analyzer" / "sessions"
STORAGE_DIR.mkdir(parents=True, exist_ok=True)

def save_session(pm_name, spec_title, spec_text, initial_analysis, lifecycle_stage,
                 conversation_history, final_summary, updated_spec=None):
    """Save the interview session in a dedicated folder with all files."""
    try:
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        session_id = hashlib.md5(f"{pm_name}{spec_title}{timestamp}".encode()).hexdigest()[:8]
        folder_name = f"{timestamp}_{session_id}_{spec_title.replace(' ', '_')[:30]}"

        # Create session folder
        session_folder = STORAGE_DIR / folder_name
        session_folder.mkdir(parents=True, exist_ok=True)

        # Save session data as JSON
        session_data = {
            "timestamp": timestamp,
            "pm_name": pm_name,
            "spec_title": spec_title,
            "spec_text": spec_text,
            "initial_analysis": initial_analysis,
            "lifecycle_stage": lifecycle_stage,
            "conversation_history": conversation_history,
            "final_summary": final_summary,
        }

        with open(session_folder / "session.json", 'w') as f:
            json.dump(session_data, f, indent=2)

        # Save original spec
        with open(session_folder / "original_spec.md", 'w') as f:
            f.write(f"# {spec_title}\n\n")
            f.write(f"**Date**: {datetime.now().strftime('%B %d, %Y')}\n\n")
            f.write("---\n\n")
            f.write(spec_text)

        # Save updated spec if available
        if updated_spec:
            with open(session_folder / "updated_spec.md", 'w') as f:
                f.write(updated_spec)

        # Save interview transcript
        with open(session_folder / "transcript.md", 'w') as f:
            f.write(f"# Review Transcript: {spec_title}\n\n")
            f.write(f"**Date**: {datetime.now().strftime('%B %d, %Y at %I:%M %p')}\n")
            f.write(f"**Stage**: {lifecycle_stage}\n\n")
            f.write("---\n\n")

            for entry in conversation_history:
                if entry['role'] == 'coach':
                    f.write(f"## Coach: {entry.get('topic', 'General')}\n\n")
                    f.write(f"{entry['content']}\n\n")
                    if entry.get('why_asking'):
                        f.write(f"*Why asking: {entry['why_asking']}*\n\n")
                else:
                    f.write(f"**Your Answer:**\n\n")
                    f.write(f"{entry['content']}\n\n")
                    f.write("---\n\n")

            if final_summary:
                f.write(f"\n## Final Summary\n\n")
                f.write(f"{final_summary.get('summary', '')}\n\n")

                if final_summary.get('key_insights'):
                    f.write("### Key Insights\n\n")
                    for insight in final_summary['key_insights']:
                        f.write(f"- {insight}\n")
                    f.write("\n")

                if final_summary.get('remaining_gaps'):
                    f.write("### Remaining Gaps\n\n")
                    for gap in final_summary['remaining_gaps']:
                        f.write(f"- {gap}\n")

        return session_folder
    except Exception as e:
        st.error(f"Error saving session: {str(e)}")
        return None

def load_past_sessions():
    """Load all past interview sessions from their folders."""
    try:
        sessions = []
        for folder in sorted(STORAGE_DIR.glob("*"), reverse=True):
            if folder.is_dir():
                try:
                    session_file = folder / "session.json"
                    if session_file.exists():
                        with open(session_file, 'r') as f:
                            data = json.load(f)
                            data['folder_path'] = str(folder)
                            data['folder_name'] = folder.name
                            sessions.append(data)
                except:
                    continue
        return sessions
    except Exception as e:
        return []

# Extract text from Word documents
def extract_text_from_docx(file_bytes):
    """Extract text from a .docx file"""
    if not DOCX_AVAILABLE:
        return None

    try:
        doc = Document(io.BytesIO(file_bytes))
        full_text = []

        # Extract text from paragraphs
        for para in doc.paragraphs:
            if para.text.strip():
                full_text.append(para.text)

        # Extract text from tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if cell.text.strip():
                        full_text.append(cell.text)

        return '\n\n'.join(full_text)
    except Exception as e:
        return None

if not CLAUDE_CLI:
    st.error("❌ Claude Code CLI not found! Please install Claude Code first: https://claude.com/claude-code")
    st.stop()

# Interview topics - what to ask about
INTERVIEW_TOPICS = {
    "Problem Definition": [
        "What is the root cause vs symptom?",
        "How do you know this is the most important problem?",
        "What happens if we do nothing?",
        "How was this problem validated with users?"
    ],
    "User Understanding": [
        "Who are the edge case users?",
        "What's the job-to-be-done?",
        "How do power users differ from casual users?",
        "What constraints do users operate under?"
    ],
    "Solution Rationale": [
        "What alternatives were considered?",
        "Why is this the simplest viable solution?",
        "What trade-offs are you accepting?",
        "Build vs buy vs partner - why this choice?"
    ],
    "Success Metrics": [
        "What are the leading indicators?",
        "What's the baseline and target?",
        "How will you measure this in practice?",
        "What would make you change course?"
    ],
    "Technical Feasibility": [
        "What are the performance implications at scale?",
        "What dependencies exist?",
        "What's the security model?",
        "What could break in production?"
    ],
    "Assumptions & Risks": [
        "What assumptions are you most uncertain about?",
        "What's the rollback plan?",
        "What could make this fail spectacularly?",
        "Which assumptions need validation first?"
    ],
    "Stakeholder Alignment": [
        "Who might object and why?",
        "What cross-functional dependencies exist?",
        "How will you communicate this?",
        "Who has veto power?"
    ],
    "Scope & Trade-offs": [
        "What are you explicitly NOT building?",
        "Why is this the MVP?",
        "What will you add in phase 2?",
        "What scope creep concerns exist?"
    ]
}

def initial_spec_analysis(spec_text, pm_name, spec_title):
    """
    Analyze the spec initially - rate it and compare to PRD templates.
    """
    prompt = f"""You are analyzing a product spec to provide an initial assessment.

SPEC TO ANALYZE:
Title: {spec_title}
PM: {pm_name}

{spec_text}

INSTRUCTIONS:
Analyze this spec against common PRD/spec templates and best practices.

Evaluate across these 8 dimensions (score 1-5 each):
1. Problem Definition - Root cause identified? User pain quantified?
2. User Understanding - Target users clear? Jobs-to-be-done defined?
3. Solution Rationale - Alternatives considered? Trade-offs documented?
4. Success Metrics - Measurable KPIs? Baselines and targets?
5. Technical Feasibility - Architecture outlined? Dependencies identified?
6. Assumptions & Risks - Key assumptions listed? Risk mitigation planned?
7. Stakeholder Alignment - Stakeholders identified? Concerns addressed?
8. Scope & Trade-offs - MVP clear? Out-of-scope items listed?

Also compare to common PRD templates:
- Amazon-style Working Backwards (PR/FAQ)
- Shape Up pitch format
- Traditional PRD structure
- One-pager format

Respond in JSON format:
{{
    "overall_score": <number 1-5>,
    "overall_assessment": "2-3 sentence high-level summary of spec quality",
    "dimension_scores": {{
        "Problem Definition": <1-5>,
        "User Understanding": <1-5>,
        "Solution Rationale": <1-5>,
        "Success Metrics": <1-5>,
        "Technical Feasibility": <1-5>,
        "Assumptions & Risks": <1-5>,
        "Stakeholder Alignment": <1-5>,
        "Scope & Trade-offs": <1-5>
    }},
    "template_completeness": "Brief assessment of how complete this is compared to standard PRD templates",
    "strengths": ["strength 1", "strength 2"],
    "critical_gaps": ["gap 1", "gap 2", "gap 3"],
    "readiness_level": "early_draft|developing|nearly_complete"
}}

Be honest and specific in your assessment."""

    try:
        result = subprocess.run(
            [CLAUDE_CLI, '--print', prompt],
            capture_output=True,
            text=True,
            timeout=60
        )

        if result.returncode != 0:
            return None

        response = result.stdout

        # Parse JSON
        if '```json' in response:
            response = response.split('```json')[1].split('```')[0]
        elif '```' in response:
            response = response.split('```')[1].split('```')[0]

        json_start = response.find('{')
        json_end = response.rfind('}') + 1

        if json_start != -1 and json_end > json_start:
            json_str = response[json_start:json_end].strip()
            return json.loads(json_str)

        return None

    except Exception as e:
        st.error(f"Error in initial analysis: {str(e)}")
        return None

def call_claude_interview(conversation_history, spec_text, pm_name, spec_title, stage=None, initial_analysis=None):
    """
    Call Claude Code to conduct an interview about the spec.
    Uses the conversation history to ask progressively deeper questions.
    Adjusts coaching based on the PM's stated lifecycle stage.
    """

    # Build conversation context
    conversation_context = "\n".join([
        f"{'PM' if entry['role'] == 'pm' else 'Coach'}: {entry['content']}"
        for entry in conversation_history
    ])

    is_first_question = len(conversation_history) == 0

    # Build stage-specific coaching context
    stage_context = ""
    if stage:
        if stage == "early":
            stage_context = """
LIFECYCLE STAGE: Early Draft
- Focus on problem clarity and user understanding first
- It's okay if solution details are sparse
- Push for clarity on the "why" before the "how"
- Be encouraging but probe deeply on problem/user fit"""
        elif stage == "middle":
            stage_context = """
LIFECYCLE STAGE: Developing/Middle Stage
- Expect problem and users to be defined
- Push on solution rationale and alternatives considered
- Challenge technical feasibility and success metrics
- Be direct about gaps that should be filled by now"""
        elif stage == "final":
            stage_context = """
LIFECYCLE STAGE: Final Review
- Expect comprehensive coverage of all 8 dimensions
- Be rigorous - this should be stakeholder-ready
- Focus on edge cases, risks, and what could go wrong
- Challenge assumptions that haven't been validated"""

    analysis_context = ""
    if initial_analysis:
        analysis_context = f"""
INITIAL ANALYSIS RESULTS:
Overall Score: {initial_analysis.get('overall_score', 'N/A')}/5
Critical Gaps: {', '.join(initial_analysis.get('critical_gaps', []))}
Template Completeness: {initial_analysis.get('template_completeness', 'N/A')}
"""

    if is_first_question:
        instructions = f"""{stage_context}
{analysis_context}

INSTRUCTIONS:
1. Acknowledge their stage and initial analysis results briefly
2. Ask ONE specific, non-obvious question about the MOST CRITICAL gap for their stage
3. Be appropriate for their lifecycle stage (early = supportive, final = rigorous)

Respond in JSON format:
{{
    "question": "Your question here (tailored to their stage)",
    "topic": "Which topic area (e.g., Problem Definition, Solution Rationale)",
    "why_asking": "Brief explanation of what you're trying to understand",
    "done": false
}}"""
    else:
        instructions = f"""{stage_context}

INSTRUCTIONS:
1. Ask ONE non-obvious question at a time about the spec
2. Focus on areas that are unclear, missing, or need deeper thinking
3. Go progressively deeper based on previous answers
4. Adjust your questioning style to their lifecycle stage (early/middle/final)
5. Topics to explore: Problem clarity, User understanding, Solution rationale, Success metrics, Technical feasibility, Assumptions/risks, Stakeholders, Scope/trade-offs

Your question should:
- Challenge an assumption
- Expose a gap
- Force deeper thinking
- Be specific to this spec
- Be appropriate for their stage

Respond in JSON format:"""

    prompt = f"""You are conducting a Socratic interview about a product spec.

SPEC BEING REVIEWED:
Title: {spec_title}
PM: {pm_name}

{spec_text}

INTERVIEW SO FAR:
{conversation_context if conversation_context else "Just starting - this is your first question."}

{instructions}
{{
    "question": "The specific question to ask",
    "topic": "Which topic area (e.g., Problem Definition, Solution Rationale)",
    "why_asking": "Brief explanation of what you're trying to understand",
    "done": false
}}

If you have enough information OR the user requested to finish, set "done": true and include:
{{
    "done": true,
    "summary": "1-2 sentence summary",
    "key_insights": ["top 3 insights from discussion"],
    "remaining_gaps": ["critical gaps still to address"] or []
}}

Keep the summary concise. Focus on actionable insights."""

    try:
        result = subprocess.run(
            [CLAUDE_CLI, '--print', prompt],
            capture_output=True,
            text=True,
            timeout=60
        )

        if result.returncode != 0:
            raise Exception(f"Claude Code error: {result.stderr}")

        response = result.stdout

        # Parse JSON
        if '```json' in response:
            response = response.split('```json')[1].split('```')[0]
        elif '```' in response:
            response = response.split('```')[1].split('```')[0]

        json_start = response.find('{')
        json_end = response.rfind('}') + 1

        if json_start != -1 and json_end > json_start:
            json_str = response[json_start:json_end].strip()
            return json.loads(json_str)

        return None

    except Exception as e:
        st.error(f"Error in interview: {str(e)}")
        return None

def generate_updated_spec(original_spec, conversation_history, summary, pm_name, spec_title):
    """
    Generate an updated spec incorporating insights from the Socratic interview.
    """
    # Build conversation summary
    qa_pairs = []
    for i, entry in enumerate(conversation_history):
        if entry['role'] == 'coach':
            question = entry['content']
            # Find the next PM response
            if i + 1 < len(conversation_history) and conversation_history[i+1]['role'] == 'pm':
                answer = conversation_history[i+1]['content']
                if answer != "[Skipped]":
                    qa_pairs.append(f"Q: {question}\nA: {answer}")

    conversation_summary = "\n\n".join(qa_pairs)

    insights_text = "\n".join([f"- {insight}" for insight in summary.get('key_insights', [])])
    gaps_text = "\n".join([f"- {gap}" for gap in summary.get('remaining_gaps', [])])

    prompt = f"""You are helping a PM update their product spec based on a Socratic interview.

ORIGINAL SPEC:
{original_spec}

INTERVIEW INSIGHTS:
{conversation_summary}

KEY INSIGHTS IDENTIFIED:
{insights_text}

REMAINING GAPS TO ADDRESS:
{gaps_text}

INSTRUCTIONS:
1. Take the original spec and improve it by incorporating the insights from the interview
2. Address the key insights discovered
3. Fill in the remaining gaps where possible
4. Keep the PM's voice and writing style
5. Use markdown formatting for structure
6. DO NOT remove existing good content - enhance it

Generate an updated, improved version of the spec that incorporates these learnings.

Format as a proper product spec with:
- Clear sections and headers
- The PM's insights integrated naturally
- Gaps addressed with thoughtful content
- Professional markdown formatting

Output the complete updated spec."""

    try:
        result = subprocess.run(
            [CLAUDE_CLI, '--print', prompt],
            capture_output=True,
            text=True,
            timeout=120
        )

        if result.returncode != 0:
            st.error(f"Error generating updated spec: {result.stderr}")
            return None

        return result.stdout.strip()

    except Exception as e:
        st.error(f"Error in spec generation: {str(e)}")
        return None

def refine_spec_section(current_spec, refinement_request, spec_title):
    """
    Refine a specific part of the spec based on PM's request.
    """
    prompt = f"""You are helping a PM refine their product spec.

CURRENT SPEC:
{current_spec}

PM'S REFINEMENT REQUEST:
{refinement_request}

INSTRUCTIONS:
1. Make the requested improvements to the spec
2. Keep everything else intact
3. Maintain the PM's voice and style
4. Use markdown formatting
5. Be thorough and specific

Output the complete updated spec with the refinement applied."""

    try:
        result = subprocess.run(
            [CLAUDE_CLI, '--print', prompt],
            capture_output=True,
            text=True,
            timeout=90
        )

        if result.returncode != 0:
            st.error(f"Error refining spec: {result.stderr}")
            return None

        return result.stdout.strip()

    except Exception as e:
        st.error(f"Error in refinement: {str(e)}")
        return None

# Session state initialization
if 'spec_text' not in st.session_state:
    st.session_state.spec_text = ""
if 'pm_name' not in st.session_state:
    st.session_state.pm_name = "PM"  # Default value
if 'spec_title' not in st.session_state:
    st.session_state.spec_title = ""
if 'interview_started' not in st.session_state:
    st.session_state.interview_started = False
if 'conversation_history' not in st.session_state:
    st.session_state.conversation_history = []
if 'interview_complete' not in st.session_state:
    st.session_state.interview_complete = False
if 'final_summary' not in st.session_state:
    st.session_state.final_summary = None
if 'initial_analysis' not in st.session_state:
    st.session_state.initial_analysis = None
if 'lifecycle_stage' not in st.session_state:
    st.session_state.lifecycle_stage = None
if 'analysis_complete' not in st.session_state:
    st.session_state.analysis_complete = False
if 'viewing_session' not in st.session_state:
    st.session_state.viewing_session = None
if 'session_saved' not in st.session_state:
    st.session_state.session_saved = False
if 'review_mode' not in st.session_state:
    st.session_state.review_mode = False
if 'editing_answer' not in st.session_state:
    st.session_state.editing_answer = None
if 'wants_to_finish' not in st.session_state:
    st.session_state.wants_to_finish = False

# Main UI
st.title("PM Spec Analyzer")
st.markdown("<p style='font-size: 16px; color: var(--color-text-muted); margin-bottom: 32px;'>One question at a time. I'll challenge your thinking, expose gaps, and help you develop stronger product specs.</p>", unsafe_allow_html=True)

# Check if viewing a past session
if st.session_state.viewing_session:
    session = st.session_state.viewing_session

    st.header(f"📚 Past Session: {session['spec_title']}")
    st.caption(f"{datetime.strptime(session['timestamp'], '%Y%m%d_%H%M%S').strftime('%B %d, %Y at %I:%M %p')}")

    if st.button("← Back to Current Session"):
        st.session_state.viewing_session = None
        st.rerun()

    st.divider()

    # Show initial analysis
    if session.get('initial_analysis'):
        st.markdown("### 📊 Initial Analysis")
        analysis = session['initial_analysis']
        col1, col2 = st.columns([1, 3])
        with col1:
            st.metric("Score", f"{analysis.get('overall_score', 'N/A')}/5")
        with col2:
            st.info(analysis.get('overall_assessment', ''))

        with st.expander("View Full Analysis"):
            st.json(analysis)

    # Show lifecycle stage
    if session.get('lifecycle_stage'):
        stage_map = {
            'early': '📝 Early Draft',
            'middle': '🔨 Developing',
            'final': '✨ Final Review'
        }
        st.markdown(f"**Stage**: {stage_map.get(session['lifecycle_stage'], session['lifecycle_stage'])}")

    # Show conversation history
    st.markdown("### 💬 Interview Transcript")
    for entry in session.get('conversation_history', []):
        if entry['role'] == 'coach':
            with st.chat_message("assistant", avatar="🤔"):
                st.markdown(f"**Topic**: {entry.get('topic', 'General')}")
                st.markdown(entry['content'])
                if entry.get('why_asking'):
                    st.caption(f"💡 {entry['why_asking']}")
        else:
            with st.chat_message("user", avatar="👤"):
                st.markdown(entry['content'])

    # Show final summary
    if session.get('final_summary'):
        st.markdown("### 📊 Final Summary")
        summary = session['final_summary']
        st.info(summary.get('summary', ''))

        if summary.get('key_insights'):
            st.markdown("**💡 Key Insights**")
            for insight in summary['key_insights']:
                st.markdown(f"- {insight}")

        if summary.get('remaining_gaps'):
            st.markdown("**⚠️ Remaining Gaps**")
            for gap in summary['remaining_gaps']:
                st.markdown(f"- {gap}")

    # Show updated spec if available
    if session.get('updated_spec'):
        st.markdown("### 📝 Updated Spec")
        with st.expander("View Updated Spec"):
            st.markdown(session['updated_spec'])

    st.stop()  # Stop here when viewing past session

# Sidebar
with st.sidebar:
    st.header("How It Works")
    st.markdown("""
    **Socratic Interview Approach**

    - 🤔 **One question at a time**
    - 🎯 **Progressive depth** based on your answers
    - 💡 **Challenge assumptions** and expose gaps
    - 🔄 **Follow-up questions** to go deeper
    - ✅ **Complete when spec is ready**

    **Time**: 10-20 minutes
    """)

    st.divider()
    st.markdown(f"**✅ Claude Code**: Found")

    if DOCX_AVAILABLE:
        st.markdown("**✅ Word Support**: Enabled (.docx)")
    else:
        st.markdown("**⚠️ Word Support**: Not available")
        st.caption("Run: `pipx inject streamlit python-docx`")

    st.markdown("**📄 Formats**: .txt, .md" + (", .docx" if DOCX_AVAILABLE else ""))

    st.divider()

    # Past sessions review
    with st.expander("📚 Review Past Sessions", expanded=False):
        past_sessions = load_past_sessions()

        if past_sessions:
            st.markdown(f"**{len(past_sessions)} saved sessions**")

            for session in past_sessions[:10]:  # Show last 10
                timestamp = datetime.strptime(session['timestamp'], "%Y%m%d_%H%M%S").strftime("%b %d, %Y %I:%M %p")

                with st.container():
                    col1, col2 = st.columns([3, 1])
                    with col1:
                        st.markdown(f"**{session['spec_title']}**")
                        st.caption(timestamp)
                    with col2:
                        if st.button("View", key=f"view_{session['filename']}"):
                            st.session_state.viewing_session = session
                    st.markdown("---")
        else:
            st.info("No past sessions yet. Complete an interview to save your first session!")

    if st.session_state.interview_started:
        st.markdown(f"**📊 Questions Asked**: {len([e for e in st.session_state.conversation_history if e['role'] == 'coach'])}")
        st.markdown(f"**💬 Your Responses**: {len([e for e in st.session_state.conversation_history if e['role'] == 'pm'])}")

# Main content
if not st.session_state.interview_started:
    # Initial setup
    st.header("Upload Your Spec")

    spec_input_method = st.radio("Input Method", ["Upload File", "Paste Text"], horizontal=True)

    spec_text = ""
    uploaded_filename = ""

    if spec_input_method == "Upload File":
        file_types = ['txt', 'md']
        if DOCX_AVAILABLE:
            file_types.append('docx')

        uploaded_file = st.file_uploader(
            "Choose file",
            type=file_types,
            help="Supports .txt, .md, and .docx files",
            label_visibility="collapsed"
        )

        if uploaded_file:
            # Auto-populate title from filename (without extension)
            uploaded_filename = uploaded_file.name.rsplit('.', 1)[0]
            file_extension = uploaded_file.name.split('.')[-1].lower()

            try:
                if file_extension == 'docx':
                    if not DOCX_AVAILABLE:
                        st.error("❌ Word document support not available. Please install python-docx.")
                        st.info("Run: `pipx inject streamlit python-docx` then restart")
                        spec_text = ""
                    else:
                        # Extract text from Word document
                        file_bytes = uploaded_file.read()
                        spec_text = extract_text_from_docx(file_bytes)

                        if spec_text and len(spec_text.strip()) > 0:
                            st.success(f"✅ Loaded {len(spec_text)} characters from {uploaded_file.name}")
                            with st.expander("Preview Document", expanded=False):
                                st.text_area("Content", spec_text, height=300, disabled=True)
                        else:
                            st.warning("⚠️ Word document appears to be empty or only contains images")
                            st.info("Try copy-pasting the content instead")
                            spec_text = ""

                elif file_extension in ['txt', 'md']:
                    # Read text/markdown files
                    spec_text = uploaded_file.read().decode('utf-8')
                    if spec_text and len(spec_text.strip()) > 0:
                        st.success(f"✅ Loaded {len(spec_text)} characters from {uploaded_file.name}")
                        with st.expander("Preview Document", expanded=False):
                            st.text_area("Content", spec_text, height=300, disabled=True)
                    else:
                        st.warning("⚠️ File appears to be empty")
                        spec_text = ""
                else:
                    st.error(f"❌ Unsupported file type: .{file_extension}")
                    spec_text = ""

            except Exception as e:
                st.error(f"❌ Error reading file: {str(e)}")
                spec_text = ""
    else:
        spec_text = st.text_area(
            "Paste Your Spec Here",
            value=st.session_state.spec_text,
            height=400,
            placeholder="Paste your product spec here...",
            label_visibility="collapsed"
        )

    # Spec title - auto-populated from filename or editable
    spec_title = st.text_input(
        "Spec Title",
        value=uploaded_filename if uploaded_filename else st.session_state.spec_title,
        placeholder="Enter spec title",
        key="spec_title_input"
    )

    st.markdown("<br>", unsafe_allow_html=True)

    if st.button("Start Review", type="primary", disabled=not (spec_title and spec_text), use_container_width=True):
        st.session_state.pm_name = "PM"  # Default name
        st.session_state.spec_title = spec_title
        st.session_state.spec_text = spec_text
        st.session_state.interview_started = True
        st.session_state.conversation_history = []
        st.session_state.interview_complete = False
        st.session_state.needs_first_question = True
        st.rerun()

else:
    # Interview in progress
    st.markdown(f"## Reviewing: {st.session_state.spec_title}")

    # Step 1: Run initial analysis if not done
    if not st.session_state.analysis_complete:
        loading_placeholder = st.empty()
        with loading_placeholder.container():
            show_loading_with_quote("🔍 Analyzing your spec against PRD best practices")

        analysis = initial_spec_analysis(
            st.session_state.spec_text,
            st.session_state.pm_name,
            st.session_state.spec_title
        )
        if analysis:
            st.session_state.initial_analysis = analysis
            st.session_state.analysis_complete = True
            loading_placeholder.empty()
            st.rerun()

    # Step 2: Show analysis and ask for stage
    if st.session_state.analysis_complete and not st.session_state.lifecycle_stage:
        analysis = st.session_state.initial_analysis

        st.markdown("## Initial Assessment")

        # Minimal score card - Linear style
        score = analysis.get('overall_score', 0)
        score_html = f"""
        <div class="score-card">
            <p class="score-label">Overall Score</p>
            <p class="score-number">{score}<span style="font-size: 36px; opacity: 0.4; font-weight: 400;">/5</span></p>
            <p class="score-label">Across 8 Key Dimensions</p>
        </div>
        """
        st.markdown(score_html, unsafe_allow_html=True)

        st.markdown(f"<p style='margin: 24px 0; font-size: 15px; line-height: 1.6; color: var(--color-text-muted);'>{analysis.get('overall_assessment', '')}</p>", unsafe_allow_html=True)

        # Dimension scores in expandable section
        with st.expander("Dimension Breakdown", expanded=False):
            dim_scores = analysis.get('dimension_scores', {})

            # Display dimensions as clean cards
            for dim, score in dim_scores.items():
                dim_html = f"""
                <div class="dimension-card">
                    <div style="display: flex; justify-content: space-between; align-items: center;">
                        <span style="font-weight: 500; font-size: 14px; color: var(--color-text);">{dim}</span>
                        <span style="font-size: 14px; font-weight: 600; color: var(--color-text-muted);">{score}/5</span>
                    </div>
                </div>
                """
                st.markdown(dim_html, unsafe_allow_html=True)

        # Strengths and gaps
        st.markdown("### Strengths")
        for strength in analysis.get('strengths', []):
            st.markdown(f"<p style='font-size: 14px; color: var(--color-text-muted); margin: 8px 0;'>✓ {strength}</p>", unsafe_allow_html=True)

        st.markdown("### Critical Gaps")
        for gap in analysis.get('critical_gaps', []):
            st.markdown(f"<p style='font-size: 14px; color: var(--color-text-muted); margin: 8px 0;'>• {gap}</p>", unsafe_allow_html=True)

        st.divider()

        # Ask for lifecycle stage - Clean survey style
        st.markdown("---")
        st.markdown("## What Stage Are You At?")
        st.markdown("<p style='font-size: 15px; color: var(--color-text-muted); margin-bottom: 24px;'>Help me coach you appropriately based on where you are in your spec development</p>", unsafe_allow_html=True)

        stage = st.radio(
            "Select your current stage:",
            options=[
                ("early", "Early Draft - Just starting, brainstorming ideas"),
                ("middle", "Developing - Fleshing out details, mid-way through"),
                ("final", "Final Review - Nearly complete, ready for stakeholders")
            ],
            format_func=lambda x: x[1],
            key="stage_selector",
            label_visibility="collapsed"
        )

        st.markdown("<br>", unsafe_allow_html=True)

        if st.button("Continue", type="primary"):
            st.session_state.lifecycle_stage = stage[0]
            st.session_state.needs_first_question = True
            st.rerun()

    # Step 3: Auto-generate first question with stage context
    elif st.session_state.get('needs_first_question', False):
        loading_placeholder = st.empty()
        with loading_placeholder.container():
            show_loading_with_quote("🤔 Coach is preparing your first question")

        result = call_claude_interview(
            [],  # Empty history for first question
            st.session_state.spec_text,
            st.session_state.pm_name,
            st.session_state.spec_title,
            st.session_state.lifecycle_stage,
            st.session_state.initial_analysis
        )

        if result and not result.get('done'):
            # Add first question
            st.session_state.conversation_history.append({
                'role': 'coach',
                'content': result.get('question', ''),
                'topic': result.get('topic', ''),
                'why_asking': result.get('why_asking', ''),
                'is_first': True
            })
            st.session_state.needs_first_question = False
            loading_placeholder.empty()
            st.rerun()

    # Show conversation history (survey UX - collapsed by default)
    if st.session_state.conversation_history and not st.session_state.interview_complete:
        # Show progress indicator
        questions_asked = len([e for e in st.session_state.conversation_history if e['role'] == 'coach'])
        estimated_total = 8  # Roughly one per dimension
        show_progress(questions_asked, estimated_total)

        # Show previous Q&A in collapsed expander with edit option
        if questions_asked > 1:
            with st.expander(f"View Previous Questions ({questions_asked - 1} answered)", expanded=False):
                # Show all but the last question/answer pair
                qa_pairs = []
                for i in range(0, len(st.session_state.conversation_history) - 2, 2):
                    if i+1 < len(st.session_state.conversation_history):
                        qa_pairs.append((i, st.session_state.conversation_history[i], st.session_state.conversation_history[i+1]))

                for idx, (orig_idx, q_entry, a_entry) in enumerate(qa_pairs):
                    topic_html = f"""
                    <div class="topic-badge">{q_entry.get('topic', 'General')}</div>
                    """
                    st.markdown(topic_html, unsafe_allow_html=True)
                    st.markdown(f"**Q:** {q_entry['content']}")

                    # Show answer with edit button
                    col1, col2 = st.columns([4, 1])
                    with col1:
                        if a_entry['content'] != "[Skipped]":
                            st.markdown(f"**A:** {a_entry['content']}")
                        else:
                            st.markdown("*[Skipped]*")
                    with col2:
                        if st.button("Edit", key=f"edit_{orig_idx}"):
                            st.session_state.editing_answer = orig_idx + 1  # Index of the answer
                            st.rerun()

                    st.markdown("---")

    # Handle editing previous answer
    if st.session_state.editing_answer is not None:
        edit_idx = st.session_state.editing_answer
        if edit_idx < len(st.session_state.conversation_history):
            q_entry = st.session_state.conversation_history[edit_idx - 1]
            a_entry = st.session_state.conversation_history[edit_idx]

            st.markdown("## Edit Your Answer")

            topic_html = f"""
            <div class="topic-badge">{q_entry.get('topic', 'General')}</div>
            """
            st.markdown(topic_html, unsafe_allow_html=True)

            question_html = f"""
            <div class="question-display">
                {q_entry['content']}
            </div>
            """
            st.markdown(question_html, unsafe_allow_html=True)

            current_answer = a_entry['content'] if a_entry['content'] != "[Skipped]" else ""
            new_answer = st.text_area(
                "Update your answer",
                value=current_answer,
                height=120,
                key="edit_answer_input"
            )

            col1, col2 = st.columns(2)
            with col1:
                if st.button("Save Changes", type="primary", use_container_width=True):
                    st.session_state.conversation_history[edit_idx]['content'] = new_answer if new_answer else "[Skipped]"
                    st.session_state.editing_answer = None
                    st.rerun()
            with col2:
                if st.button("Cancel", use_container_width=True):
                    st.session_state.editing_answer = None
                    st.rerun()

            st.stop()  # Don't show rest of UI while editing

    # Review mode - show all Q&A before submitting
    if st.session_state.review_mode and not st.session_state.interview_complete:
        st.markdown("## Review Your Answers")
        st.markdown("<p style='font-size: 15px; color: var(--color-text-muted); margin-bottom: 24px;'>Review all your answers below. You can edit any answer before submitting.</p>", unsafe_allow_html=True)

        # Show all Q&A pairs
        qa_pairs = []
        for i in range(0, len(st.session_state.conversation_history), 2):
            if i+1 < len(st.session_state.conversation_history):
                qa_pairs.append((i, st.session_state.conversation_history[i], st.session_state.conversation_history[i+1]))

        for idx, (orig_idx, q_entry, a_entry) in enumerate(qa_pairs):
            with st.container():
                st.markdown(f"### Question {idx + 1}")
                topic_html = f"""
                <div class="topic-badge">{q_entry.get('topic', 'General')}</div>
                """
                st.markdown(topic_html, unsafe_allow_html=True)
                st.markdown(f"**{q_entry['content']}**")

                col1, col2 = st.columns([5, 1])
                with col1:
                    if a_entry['content'] != "[Skipped]":
                        st.markdown(f"<div class='question-display'>{a_entry['content']}</div>", unsafe_allow_html=True)
                    else:
                        st.markdown("<div class='question-display'><em>[Skipped]</em></div>", unsafe_allow_html=True)
                with col2:
                    if st.button("Edit", key=f"review_edit_{orig_idx}", use_container_width=True):
                        st.session_state.editing_answer = orig_idx + 1
                        st.rerun()

                st.markdown("---")

        # Action buttons
        st.markdown("<br>", unsafe_allow_html=True)

        col1, col2, col3 = st.columns([2, 2, 1])

        with col1:
            if st.button("Submit & Get Feedback", type="primary", use_container_width=True):
                # Mark as wanting to finish and proceed to completion
                st.session_state.wants_to_finish = True
                st.session_state.review_mode = False

                # Show detailed progress
                progress_placeholder = st.empty()

                with progress_placeholder.container():
                    st.markdown("### Analyzing Your Answers")

                    status_html = """
                    <div class="dimension-card">
                        <p style="margin: 0; font-size: 14px;">
                            Reviewing all your answers and preparing feedback...
                        </p>
                    </div>
                    """
                    st.markdown(status_html, unsafe_allow_html=True)
                    show_loading_with_quote("This may take 30-60 seconds")

                # Get final summary/coaching (in background)
                result = call_claude_interview(
                    st.session_state.conversation_history,
                    st.session_state.spec_text,
                    st.session_state.pm_name,
                    st.session_state.spec_title,
                    st.session_state.lifecycle_stage,
                    st.session_state.initial_analysis
                )

                progress_placeholder.empty()

                if result and result.get('done'):
                    st.session_state.interview_complete = True
                    st.session_state.final_summary = result

                    # Show completion message
                    st.success("✓ Analysis complete!")
                    time.sleep(0.5)
                    st.rerun()

        with col2:
            if st.button("Continue Asking", use_container_width=True):
                st.session_state.review_mode = False
                st.session_state.needs_next_question = True
                st.rerun()

        with col3:
            if st.button("Cancel", use_container_width=True):
                st.session_state.review_mode = False
                st.rerun()

        st.stop()  # Don't show rest of UI while in review mode

    # Auto-advance logic - fetch next question automatically after answer/skip
    if st.session_state.get('needs_next_question', False):
        loading_placeholder = st.empty()
        with loading_placeholder.container():
            show_loading_with_quote("Thinking of next question")

        result = call_claude_interview(
            st.session_state.conversation_history,
            st.session_state.spec_text,
            st.session_state.pm_name,
            st.session_state.spec_title,
            st.session_state.lifecycle_stage,
            st.session_state.initial_analysis
        )

        if result:
            loading_placeholder.empty()
            st.session_state.needs_next_question = False

            if result.get('done'):
                # Interview complete
                st.session_state.interview_complete = True
                st.session_state.final_summary = result

                # Auto-save session
                if not st.session_state.session_saved:
                    session_folder = save_session(
                        st.session_state.pm_name,
                        st.session_state.spec_title,
                        st.session_state.spec_text,
                        st.session_state.initial_analysis,
                        st.session_state.lifecycle_stage,
                        st.session_state.conversation_history,
                        result,
                        st.session_state.get('updated_spec')
                    )
                    if session_folder:
                        st.session_state.session_folder = str(session_folder)
                    st.session_state.session_saved = True

                st.rerun()
            else:
                # Add question to history
                st.session_state.conversation_history.append({
                    'role': 'coach',
                    'content': result.get('question', ''),
                    'topic': result.get('topic', ''),
                    'why_asking': result.get('why_asking', '')
                })
                st.rerun()

    # Show current question if there is one
    if st.session_state.conversation_history and st.session_state.conversation_history[-1]['role'] == 'coach' and not st.session_state.interview_complete:
        last_question = st.session_state.conversation_history[-1]

        # Show topic badge
        topic_html = f"""
        <div class="topic-badge">{last_question.get('topic', 'General')}</div>
        """
        st.markdown(topic_html, unsafe_allow_html=True)

        # Show current question prominently
        question_html = f"""
        <div class="question-display">
            {last_question['content']}
        </div>
        """
        st.markdown(question_html, unsafe_allow_html=True)

        # Response textarea
        response = st.text_area(
            "Your answer",
            height=120,
            placeholder="Type your answer here...",
            key="current_response",
            help="Be thorough - explain your thinking"
        )

        # Button row - minimal
        col1, col2, col3, col4 = st.columns([3, 1, 1, 1])

        with col1:
            if st.button("Continue", type="primary", disabled=not response, use_container_width=True):
                st.session_state.conversation_history.append({
                    'role': 'pm',
                    'content': response
                })
                # Auto-advance: immediately ask next question
                st.session_state.needs_next_question = True
                st.rerun()

        with col2:
            if st.button("Skip", use_container_width=True):
                st.session_state.conversation_history.append({
                    'role': 'pm',
                    'content': "[Skipped]"
                })
                # Auto-advance: immediately ask next question
                st.session_state.needs_next_question = True
                st.rerun()

        with col3:
            if st.button("Review", use_container_width=True):
                # Go to review mode
                st.session_state.review_mode = True
                st.rerun()

        with col4:
            if st.button("End", use_container_width=True):
                # Go directly to completion
                st.session_state.review_mode = True
                st.session_state.wants_to_finish = True
                st.rerun()

    # Check if interview is complete
    if st.session_state.interview_complete and st.session_state.final_summary:
        # Minimal celebration
        complete_html = """
        <div class="celebration">
            <div class="celebration-icon">✓</div>
            <div class="celebration-title">Review Complete</div>
            <div class="celebration-subtitle">Here's your feedback</div>
        </div>
        """
        st.markdown(complete_html, unsafe_allow_html=True)

        summary = st.session_state.final_summary

        st.markdown("## Feedback Summary")
        st.markdown(f"<p style='font-size: 15px; line-height: 1.6; color: var(--color-text-muted); margin: 16px 0;'>{summary.get('summary', '')}</p>", unsafe_allow_html=True)

        if summary.get('key_insights'):
            st.markdown("### Key Insights")
            for insight in summary['key_insights']:
                st.markdown(f"<p style='font-size: 14px; color: var(--color-text-muted); margin: 8px 0;'>• {insight}</p>", unsafe_allow_html=True)

        if summary.get('remaining_gaps'):
            st.markdown("### Remaining Gaps")
            for gap in summary['remaining_gaps']:
                st.markdown(f"<p style='font-size: 14px; color: var(--color-text-muted); margin: 8px 0;'>• {gap}</p>", unsafe_allow_html=True)
        else:
            st.markdown("<p style='font-size: 14px; color: var(--color-success); margin: 16px 0;'>✓ No major gaps remaining! Spec is ready for stakeholder review.</p>", unsafe_allow_html=True)

        st.markdown("---")

        # Co-creation section
        if 'updated_spec' not in st.session_state:
            st.session_state.updated_spec = None
        if 'cocreation_mode' not in st.session_state:
            st.session_state.cocreation_mode = False

        st.markdown("---")

        # Next steps guidance
        st.markdown("### Next Steps")
        next_steps_html = """
        <div class="dimension-card">
            <p style="margin: 8px 0; font-size: 14px; color: var(--color-text);">
                <strong>1.</strong> Review the feedback and insights above<br>
                <strong>2.</strong> Incorporate the key insights into your spec<br>
                <strong>3.</strong> Address the remaining gaps identified<br>
                <strong>4.</strong> Re-upload your updated spec for another review
            </p>
        </div>
        """
        st.markdown(next_steps_html, unsafe_allow_html=True)

        st.markdown("<br>", unsafe_allow_html=True)

        # Download and save options
        col1, col2 = st.columns(2)

        with col1:
            # Download feedback as markdown
            feedback_md = f"""# Spec Review Feedback: {st.session_state.spec_title}

**Date**: {datetime.now().strftime('%B %d, %Y at %I:%M %p')}
**Stage**: {st.session_state.lifecycle_stage}

---

## Summary

{summary.get('summary', '')}

## Key Insights

"""
            for insight in summary.get('key_insights', []):
                feedback_md += f"- {insight}\n"

            if summary.get('remaining_gaps'):
                feedback_md += "\n## Remaining Gaps to Address\n\n"
                for gap in summary['remaining_gaps']:
                    feedback_md += f"- {gap}\n"

            feedback_md += "\n---\n\n## Interview Transcript\n\n"
            for entry in st.session_state.conversation_history:
                if entry['role'] == 'coach':
                    feedback_md += f"**Coach ({entry.get('topic', 'General')}):** {entry['content']}\n\n"
                else:
                    feedback_md += f"**Your Answer:** {entry['content']}\n\n"

            st.download_button(
                label="Download Feedback",
                data=feedback_md,
                file_name=f"feedback_{st.session_state.spec_title.replace(' ', '_')}.md",
                mime="text/markdown",
                type="primary",
                use_container_width=True
            )

        with col2:
            if st.button("Start New Review", use_container_width=True):
                st.session_state.interview_started = False
                st.session_state.conversation_history = []
                st.session_state.interview_complete = False
                st.session_state.final_summary = None
                st.rerun()

        # Show session saved location
        if st.session_state.get('session_folder'):
            st.markdown("<br>", unsafe_allow_html=True)
            st.markdown(f"<p style='font-size: 13px; color: var(--color-text-subtle); text-align: center;'>Session saved to: <code>{st.session_state.session_folder}</code></p>", unsafe_allow_html=True)

# Footer
st.divider()
st.markdown("""
<div style='text-align: center; color: #666; font-size: 0.9em;'>
Interview-style spec review • One question at a time • Going deep 🎯
</div>
""", unsafe_allow_html=True)
